from docx import Document
from openpyxl import load_workbook

#importamos el hoja de calculo y activamos la hoja
wb=load_workbook(filename="new_xlsx.xlsx")
ws=wb['E0.30']

#creamos objeto documento
document = Document()

#agregamos titulo
document.add_heading('Reporte de E0.30', 0)

#parrafo
p = document.add_paragraph('Se tiene que conocer la cortante vasal ')
#cursos
p.add_run('P=ZUCS/R').bold = True
p.add_run('con la condicion de que podamos saber cuanto de resistencia absorve cada columna o placa y ello verificarlo con la ')
p.add_run('norma E0.30.').italic = True

#encabezado 1
document.add_heading('Z(Zona sismica)', level=1)
document.add_paragraph('Zona 1', style='List Bullet')
document.add_paragraph('Zona 2', style='List Bullet')
document.add_paragraph('Zona 3', style='List Bullet')
document.add_paragraph('Zona 4', style='List Bullet')
document.add_paragraph(ws['C1'].value,style='Normal')

#encabezado 2
document.add_heading('U', level=1)
document.add_paragraph('tipo de uso de la edificacion',style='Normal')
document.add_paragraph(str(ws['C2'].value),style='Normal')

#encabezado 3
document.add_heading('C', level=1)
document.add_paragraph('el tipo de amplificacion',style='Normal')
document.add_paragraph(str(ws['C3'].value),style='Normal')

#encabezado 4
document.add_heading('S', level=1)
document.add_paragraph('el tipo de suelo',style='Normal')
document.add_paragraph(str(ws['C4'].value),style='Normal')

#tipos de suelo
records = (
    ('S1', '1', '1','1', '1'),
    ('S2', '1', '1','1', '1'),
    ('S3', '1', '1','1', '1'),
    ('S4', '1', '1','1', '1')
)

table = document.add_table(rows=1, cols=5)
hdr_cells = table.rows[0].cells
hdr_cells[0].text = ''
hdr_cells[1].text = 'S1'
hdr_cells[2].text = 'S2'
hdr_cells[3].text = 'S3'
hdr_cells[4].text = 'S4'
for s_0,s_1,s_2,s_3,s_4 in records:
    row_cells = table.add_row().cells
    row_cells[0].text = s_0
    row_cells[1].text = s_1
    row_cells[2].text = s_2
    row_cells[3].text = s_3
    row_cells[4].text = s_4

#encabezado 5
document.add_heading('R', level=1)
document.add_paragraph('la reduccion sismica',style='Normal')
document.add_paragraph('Ip(irregularidad en planta)', style='List Bullet')
document.add_paragraph('Ia(irregularidad en altura)', style='List Bullet')
document.add_paragraph("R0: "+str(ws['C5'].value),style='Normal')
document.add_paragraph("Ia: "+str(ws['C6'].value),style='Normal')
document.add_paragraph("Ip: "+str(ws['C7'].value),style='Normal')

document.add_page_break()
document.add_paragraph("P(ZUCS/R0*Ia*Ip): "+ws['b8'].value,style='Normal')

#guardamos el archivo
document.save('new_docx.docx')